# avery_labels.py
"""
Avery label filling helpers for .docx templates that use a 10x5 table where
columns [0,2,4] are label cells (30 labels) and columns [1,3] are gaps.

Usage:
    from avery_labels import fill_avery_30up

    labels = [
        "Cayden Schalk\n512 Crafton Ave\nFranklin, TN 37064",
        "Jane Doe\n456 Oak Ave\nNashville, TN 37201",
    ]

    fill_avery_30up(
        template_path="Avery5160AddressLabels.docx",
        output_path="filled_labels.docx",
        labels=labels,
        start_at=0,
        centered=True,
    )
"""

from __future__ import annotations

from typing import Iterable, List, Optional, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


def _clear_cell(cell) -> None:
    """Clear all paragraph text in a cell (keeps paragraph objects)."""
    for p in cell.paragraphs:
        p.clear()


def set_cell_text_multiline(cell, text: str) -> None:
    """
    Write text into a cell using multiple paragraphs (one per line).
    Note: paragraph spacing may differ between lines depending on template styles.
    """
    _clear_cell(cell)
    lines = text.split("\n") if text else []
    if not lines:
        return

    cell.paragraphs[0].add_run(lines[0])
    for line in lines[1:]:
        cell.add_paragraph(line)


def set_cell_text_centered_single_paragraph(cell, text: str) -> None:
    """
    Write text into a cell centered both vertically and horizontally.
    Uses a *single paragraph* with line breaks to avoid paragraph spacing issues.
    """
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    _clear_cell(cell)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = text.split("\n") if text else []
    if not lines:
        return

    run = p.add_run(lines[0])
    for line in lines[1:]:
        run.add_break()  # line break, not a new paragraph
        run.add_text(line)


def get_table_shape(doc: Document, table_index: int = 0) -> tuple[int, int]:
    """Return (rows, cols) for a given table index."""
    if not doc.tables:
        raise RuntimeError("No tables found in the document.")
    t = doc.tables[table_index]
    return len(t.rows), len(t.columns)


def fill_avery_30up(
    template_path: str,
    output_path: str,
    labels: Sequence[str],
    *,
    start_at: int = 0,
    centered: bool = True,
    table_index: int = 0,
    label_cols: Sequence[int] = (0, 2, 4),
    expected_shape: Optional[tuple[int, int]] = (10, 5),
    clear_unused: bool = True,
) -> None:
    """
    Fill an Avery-style 30-up label sheet represented as a 10x5 table where label
    columns are typically [0,2,4] (and [1,3] are gaps).

    Args:
        template_path: Input .docx template path.
        output_path: Output .docx path to write.
        labels: List/sequence of label strings. Use '\\n' for line breaks.
        start_at: Number of label positions to skip at the beginning (0..29).
        centered: If True, center text (single paragraph + line breaks). If False,
                  write using paragraphs (one per line).
        table_index: Which table in the document contains the labels.
        label_cols: Columns in the table that correspond to actual labels.
        expected_shape: If set, validate the table is this (rows, cols).
        clear_unused: If True, blanks all label cells before filling.

    Raises:
        RuntimeError / ValueError if the table shape or indices are unexpected.
    """
    doc = Document(template_path)

    if not doc.tables:
        raise RuntimeError("No tables found. This template might be using text boxes/shapes.")

    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"table_index out of range. Document has {len(doc.tables)} tables.")

    table = doc.tables[table_index]
    rows = len(table.rows)
    cols = len(table.columns)

    if expected_shape is not None:
        er, ec = expected_shape
        if (rows, cols) != (er, ec):
            raise RuntimeError(f"Expected table {er}x{ec}, got {rows}x{cols} (table_index={table_index}).")

    for c in label_cols:
        if c < 0 or c >= cols:
            raise ValueError(f"label_cols contains invalid column {c} for a {cols}-column table.")

    # Row-major list of label cells only
    label_cells = [table.cell(r, c) for r in range(rows) for c in label_cols]
    capacity = len(label_cells)

    if not (0 <= start_at <= capacity):
        raise ValueError(f"start_at must be between 0 and {capacity}")

    writer = set_cell_text_centered_single_paragraph if centered else set_cell_text_multiline

    if clear_unused:
        for cell in label_cells:
            writer(cell, "")

    # Fill
    max_fill = min(capacity - start_at, len(labels))
    for i in range(max_fill):
        writer(label_cells[start_at + i], labels[i])

    doc.save(output_path)
